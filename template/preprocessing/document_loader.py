"""
Carga y extrae texto de diferentes formatos de documentos
"""
from pathlib import Path
from typing import List, Dict, Any, Optional
import pypdf
from docx import Document
import openpyxl


class DocumentLoader:
    """Cargador de documentos multi-formato"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.xlsx', '.xls'}
    
    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """
        Extrae texto de un archivo PDF
        
        Args:
            file_path: Ruta al archivo PDF
        
        Returns:
            Texto extraído
        """
        text = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(page_text)
                except Exception as e:
                    print(f"Error extrayendo página {page_num + 1}: {e}")
        
        return "\n\n".join(text)
    
    @staticmethod
    def load_docx(file_path: Path) -> str:
        """
        Extrae texto de un archivo DOCX
        
        Args:
            file_path: Ruta al archivo DOCX
        
        Returns:
            Texto extraído
        """
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extraer texto de tablas también
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)
        
        return "\n\n".join(text)
    
    @staticmethod
    def load_txt(file_path: Path) -> str:
        """
        Carga archivo de texto plano
        
        Args:
            file_path: Ruta al archivo TXT
        
        Returns:
            Contenido del archivo
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def load_excel(file_path: Path) -> str:
        """
        Extrae texto de un archivo Excel
        
        Args:
            file_path: Ruta al archivo Excel
        
        Returns:
            Texto extraído de todas las hojas
        """
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"=== Hoja: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text.append(row_text)
        
        return "\n\n".join(text)
    
    @classmethod
    def load_document(cls, file_path: str | Path) -> Dict[str, Any]:
        """
        Carga un documento y extrae su texto
        
        Args:
            file_path: Ruta al archivo
        
        Returns:
            Dict con 'text', 'metadata', 'source'
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Formato no soportado: {extension}. "
                f"Soportados: {cls.SUPPORTED_EXTENSIONS}"
            )
        
        # Extraer texto según el tipo
        if extension == '.pdf':
            text = cls.load_pdf(file_path)
        elif extension == '.docx':
            text = cls.load_docx(file_path)
        elif extension in {'.txt', '.md'}:
            text = cls.load_txt(file_path)
        elif extension in {'.xlsx', '.xls'}:
            text = cls.load_excel(file_path)
        else:
            raise ValueError(f"Formato no implementado: {extension}")
        
        # Metadata básica
        metadata = {
            'source': str(file_path),
            'filename': file_path.name,
            'extension': extension,
            'size_bytes': file_path.stat().st_size
        }
        
        return {
            'text': text,
            'metadata': metadata,
            'source': str(file_path)
        }
    
    @classmethod
    def load_directory(
        cls,
        directory_path: str | Path,
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Carga todos los documentos de un directorio
        
        Args:
            directory_path: Ruta al directorio
            recursive: Buscar en subdirectorios
            extensions: Filtrar por extensiones específicas (ej: ['.pdf', '.docx'])
        
        Returns:
            Lista de documentos cargados
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directorio no encontrado: {directory_path}")
        
        if not directory_path.is_dir():
            raise ValueError(f"No es un directorio: {directory_path}")
        
        # Determinar extensiones a buscar
        search_extensions = extensions if extensions else cls.SUPPORTED_EXTENSIONS
        
        # Buscar archivos
        pattern = "**/*" if recursive else "*"
        documents = []
        
        for ext in search_extensions:
            for file_path in directory_path.glob(f"{pattern}{ext}"):
                if file_path.is_file():
                    try:
                        doc = cls.load_document(file_path)
                        documents.append(doc)
                        print(f"✓ Cargado: {file_path.name}")
                    except Exception as e:
                        print(f"✗ Error cargando {file_path.name}: {e}")
        
        return documents